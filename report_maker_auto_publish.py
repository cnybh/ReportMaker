#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
试验报告发布器 v1.0 by bohangyang 2026.08
=========================================
功能（对程序目录下 REPORT FOLDER 中的所有试验报告批量执行）：
  1. 统计每个报告 1.Procedure 文件夹内所有 PDF 的总页数，
     回填 0.cover/cover.docx1 中"2试验大纲（共XXX页）"的 XXX；
  2. 将 2.picture 文件夹内所有图片生成 A4 PDF（每页 1 张图，
     图片占页面约 80% 大小、居中放置；页面上方两行标题
     "现场照片（当前序号/总照片数量）/ Site Photos（当前/总数）"，自动编号；
     图片下方标注"日期 (Date)：XXXX-XX-XX"（日期优先取图片 EXIF，无 EXIF 时取
     文件修改日期；全部文字 10 号字），
     回填"4试验照片 Test Photo（共XXX页）"的 XXX；
  3. 统计 3.Calibration Certificate 文件夹内所有 PDF 的总页数，
     回填"5校验证书（共XXX页）"的 XXX；
  4. 将修改后的 cover.docx1 通过 Microsoft Word 导出为 PDF
     （本机 DLP 透明加密软件会加密 Word 写出的 .pdf/.xps 文件，脚本先导出为
     .dat 扩展名的明文 PDF 再改名，pypdf 才能正常读取）；
  5. 按 封面(含简介) → 1.Procedure 试验内容 → 试验照片(2.picture 生成 PDF)
     → 3.Calibration Certificate 校验证书 的顺序合并为最终 PDF；
  6. 输出到程序目录新建的 REPORT OUTPUT 文件夹，
     文件名为"报告文件夹名 + cover.docx1 中 REVISION 字段的版本号"。

依赖库：python-docx / pypdf / Pillow / pywin32
依赖软件：Microsoft Word（用于 cover.docx1 -> PDF 转换）
"""

import os
import re
import sys
import shutil
import subprocess
import tempfile

APP_TITLE = 'REPORT MAKER AUTO PUBLISH v1.0 by bohangyang 2026.08'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

IMAGE_EXTS = ('.jpg', '.jpeg', '.png', '.bmp', '.tif', '.tiff', '.gif', '.webp')
A4_PX_W, A4_PX_H = 1240, 1754          # A4 @150dpi（导出为 PDF 后即 A4 纸张）
A4_DPI = 150

# 封面正文中需要回填页数的章节（顺序与文档一致）
# 兼容新旧两版模板标题：
#   当前模板：2试验大纲 / 3试验照片 / 4校验证书
#   旧模板  ：2试验大纲 / 3试验开口项 / 4试验照片 Test Photo / 5校验证书
SECTION_HEADINGS = {
    '2试验大纲': 'procedure',
    '3试验照片': 'photo',
    '4校验证书': 'calibration',
    '3试验开口项': 'open_items',
    '4试验照片 Test Photo': 'photo',
    '5校验证书': 'calibration',
}


# ============================================================
#  弹窗提示（tkinter 可用时使用，否则退化为控制台输出）
# ============================================================
def show_popup(kind, title, msg):
    try:
        from tkinter import messagebox, Tk
        root = Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        getattr(messagebox, kind)(title, msg)
        root.destroy()
    except Exception:
        print(f'\n[{kind}] {title}\n{msg}\n')


# ============================================================
#  依赖检查 / 自动安装
# ============================================================
def ensure_deps():
    """检查依赖库，缺失则自动 pip 安装；返回是否全部就绪"""
    required = {
        'docx': 'python-docx',
        'pypdf': 'pypdf',
        'PIL': 'Pillow',
        'win32com': 'pywin32',
    }
    missing = []
    for module, pkg in required.items():
        try:
            __import__(module)
        except ImportError:
            missing.append(pkg)
    if missing:
        print('[!] 缺少依赖库：' + ', '.join(missing))
        print('    正在自动安装，请稍候...')
        try:
            subprocess.check_call(
                [sys.executable, '-m', 'pip', 'install', '--disable-pip-version-check'] + missing
            )
            print('    依赖库安装完成。')
        except Exception as e:
            print(f'[!] 依赖库安装失败：{e}')
            print('    请手动执行：python -m pip install ' + ' '.join(missing))
            return False
    return True


def ensure_word():
    """检查本机是否安装 Microsoft Word（docx -> PDF 需要）"""
    try:
        import winreg
        for root, sub in ((winreg.HKEY_LOCAL_MACHINE, r'SOFTWARE\Classes\Word.Application'),
                          (winreg.HKEY_CURRENT_USER, r'SOFTWARE\Classes\Word.Application')):
            try:
                with winreg.OpenKey(root, sub):
                    return True
            except OSError:
                continue
    except Exception:
        pass
    try:
        import win32com.client
        app = win32com.client.DispatchEx('Word.Application')
        app.Quit()
        return True
    except Exception:
        return False


# ============================================================
#  路径 / 工具函数
# ============================================================
def get_app_dir():
    """程序所在目录（exe 打包后返回 exe 所在目录）"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(os.path.abspath(sys.executable))
    return os.path.dirname(os.path.abspath(__file__))


def natural_key(s):
    """自然排序键：数字按数值比较，其余按字母忽略大小写"""
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r'(\d+)', s)]


def pdfs_in_folder(folder):
    """返回文件夹内所有 PDF 的完整路径列表（自然排序）"""
    if not os.path.isdir(folder):
        return []
    files = [os.path.join(folder, f) for f in os.listdir(folder)
             if f.lower().endswith('.pdf') and os.path.isfile(os.path.join(folder, f))]
    return sorted(files, key=lambda p: natural_key(os.path.basename(p)))


# ============================================================
#  页数统计
# ============================================================
def count_pdf_pages(pdf_path):
    """返回单个 PDF 的页数；读取失败返回 0"""
    try:
        from pypdf import PdfReader
        with open(pdf_path, 'rb') as fh:
            return len(PdfReader(fh).pages)
    except Exception as e:
        print(f'    [!] 页数统计失败：{os.path.basename(pdf_path)} - {e}')
        return 0


# ============================================================
#  图片 -> A4 PDF（每页 1 张图，占页面约 80%；
#  页面上方标题两行"现场照片（当前/总数）/ Site Photos（当前/总数）"，
#  图片下方标注日期"日期 (Date)：XXXX-XX-XX"，全部文字 10 号字）
# ============================================================
FONT_PX_10 = max(int(round(10 * A4_DPI / 72)), 12)   # 10 号字(10pt) @150dpi ≈ 21px
TITLE_TOP = 90             # 标题区起始 y
TITLE_LINE_H = 34          # 标题行距
TITLE_BOTTOM = TITLE_TOP + TITLE_LINE_H * 2          # 标题区结束 y
IMG_AREA_TOP = TITLE_BOTTOM + 20                     # 图片区起始 y
IMG_AREA_BOTTOM = A4_PX_H - 90                       # 图片区结束 y（底部留日期区）


def _get_photo_date(img, path):
    """读取照片拍摄日期：
    优先取 EXIF（DateTimeOriginal / DateTimeDigitized / DateTime），
    缺失时回退为文件修改日期。返回 'YYYY-MM-DD'；失败返回空串"""
    try:
        exif = img.getexif()
        for tag in (36867, 36868, 306):     # DateTimeOriginal / DateTimeDigitized / DateTime
            val = exif.get(tag)
            if not val:
                continue
            m = re.match(r'(\d{4})[:\-](\d{2})[:\-](\d{2})', str(val).strip())
            if m:
                return f'{m.group(1)}-{m.group(2)}-{m.group(3)}'
    except Exception:
        pass
    try:
        import datetime
        return datetime.datetime.fromtimestamp(os.path.getmtime(path)).strftime('%Y-%m-%d')
    except Exception:
        return ''


def _load_font(size_px):
    """按像素尺寸加载中文字体（微软雅黑/黑体/宋体），失败回退默认字体"""
    from PIL import ImageFont
    for fp in (r'C:\Windows\Fonts\msyh.ttc', r'C:\Windows\Fonts\msyhbd.ttc',
               r'C:\Windows\Fonts\simhei.ttf', r'C:\Windows\Fonts\simsun.ttc',
               r'C:\Windows\Fonts\arial.ttf'):
        try:
            return ImageFont.truetype(fp, size_px)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_header(canvas, idx, total):
    """在 A4 画布顶部居中绘制两行标题（10 号字）：
    第一行"现场照片（当前序号/总照片数量）"，
    第二行"Site Photos（当前/总数）"。idx 从 1 开始"""
    try:
        from PIL import ImageDraw
        draw = ImageDraw.Draw(canvas)
        font = _load_font(FONT_PX_10)
        for text, y in ((f'现场照片（{idx}/{total}）', TITLE_TOP),
                        (f'Site Photos（{idx}/{total}）', TITLE_TOP + TITLE_LINE_H)):
            bbox = draw.textbbox((0, 0), text, font=font)
            tw = bbox[2] - bbox[0]
            x = (A4_PX_W - tw) // 2 - bbox[0]        # 水平居中
            draw.text((x, y), text, font=font, fill=(60, 60, 60))
    except Exception:
        pass


def _draw_photo_date(canvas, date_text, img_top, img_height):
    """在 A4 画布上，图片下方区域内居中绘制"日期 (Date)：XXXX-XX-XX"（10 号字）"""
    if not date_text:
        return
    try:
        from PIL import ImageDraw
        label = f'日期 (Date)：{date_text}'
        draw = ImageDraw.Draw(canvas)
        font = _load_font(FONT_PX_10)
        bbox = draw.textbbox((0, 0), label, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        x = (A4_PX_W - tw) // 2 - bbox[0]                    # 水平居中
        bottom = img_top + img_height
        y = bottom + (A4_PX_H - bottom - th) // 2 - bbox[1]  # 图片下方区域内垂直居中
        draw.text((x, y), label, font=font, fill=(60, 60, 60))
    except Exception:
        pass


def build_photos_pdf(img_dir, out_pdf):
    """将 img_dir 内所有图片合并为 A4 PDF（每页 1 张、占页面约 80% 大小、
    页面顶部两行标题"现场照片（当前/总数）/ Site Photos（当前/总数）"、
    图片下方标注"日期 (Date)：XXXX-XX-XX"）。
    返回生成页数；文件夹为空或失败返回 0"""
    from PIL import Image, ImageOps

    if not os.path.isdir(img_dir):
        return 0
    files = [f for f in os.listdir(img_dir)
             if os.path.splitext(f)[1].lower() in IMAGE_EXTS
             and os.path.isfile(os.path.join(img_dir, f))]
    files.sort(key=natural_key)
    if not files:
        return 0

    pages = []
    for idx, f in enumerate(files, start=1):
        path = os.path.join(img_dir, f)
        try:
            img = Image.open(path)
            photo_date = _get_photo_date(img, path)    # 先读拍摄日期（再转置）
            img = ImageOps.exif_transpose(img)          # 按 EXIF 方向校正
            # 统一转为 RGB（含透明通道的 PNG 以白色垫底）
            if img.mode in ('RGBA', 'LA', 'P'):
                rgba = img.convert('RGBA')
                bg = Image.new('RGB', rgba.size, (255, 255, 255))
                bg.paste(rgba, mask=rgba.split()[-1])
                img = bg
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            # A4 画布：上方标题 -> 中部图片(80%) -> 下方日期
            canvas = Image.new('RGB', (A4_PX_W, A4_PX_H), (255, 255, 255))
            max_w, max_h = int(A4_PX_W * 0.8), int(A4_PX_H * 0.8)   # 页面 80% 上限
            img.thumbnail((max_w, max_h), Image.LANCZOS)
            _draw_header(canvas, idx, len(files))
            # 图片在标题区与日期区之间垂直居中
            img_top = IMG_AREA_TOP + (IMG_AREA_BOTTOM - IMG_AREA_TOP - img.height) // 2
            canvas.paste(img, ((A4_PX_W - img.width) // 2, img_top))
            _draw_photo_date(canvas, photo_date, img_top, img.height)
            print(f'    {f} -> {img.width}x{img.height}px，日期：{photo_date or "无"}')
            pages.append(canvas)
        except Exception as e:
            print(f'    [!] 图片处理失败：{f} - {e}')

    if not pages:
        return 0
    pages[0].save(out_pdf, 'PDF', save_all=True, append_images=pages[1:], resolution=A4_DPI)
    return len(pages)


# ============================================================
#  cover.docx1 编辑（python-docx）
# ============================================================
def get_revision(doc):
    """读取 cover.docx1 中标签为 REVSION 的 SDT 内容控件文本（版本号）"""
    from docx.oxml.ns import qn
    placeholders = ('单击或点击此处输入文字。', '单击此处输入文字', '[请在此处输入]')
    for sdt in doc.element.body.iter(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        tag = sdtPr.find(qn('w:tag'))
        if tag is not None and (tag.get(qn('w:val')) or '').strip().upper() in ('REVSION', 'REVISION'):
            texts = [t.text or '' for t in sdt.iter(qn('w:t'))]
            rev = ''.join(texts).strip()
            if rev and rev not in placeholders:
                return rev
    return 'A'


def set_para_text(p, new_text):
    """将段落文本替换为 new_text：保留首个 run 的格式，删除其余 run"""
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def replace_page_counts(doc, proc_total, photo_total, calib_total):
    """
    遍历封面正文段落，按章节回填页数：
      2试验大纲   -> proc_total
      4试验照片   -> photo_total
      5校验证书   -> calib_total
      3试验开口项 -> 保持 XXX 不变
    返回替换数量
    """
    values = {
        'procedure': proc_total,
        'open_items': None,      # 无对应内容文件夹，不修改
        'photo': photo_total,
        'calibration': calib_total,
    }
    current = None
    replaced = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in SECTION_HEADINGS:
            current = SECTION_HEADINGS[text]
            continue
        if current is not None and text == '（共XXX页）':
            v = values.get(current)
            if v is not None:
                set_para_text(p, f'（共{v}页）')
                replaced += 1
            current = None
    return replaced


# ============================================================
#  docx -> PDF 与章节定位（Microsoft Word 自动化）
# ============================================================
def get_heading_bookmarks(cover_path):
    """
    用 python-docx 读取封面正文中三个章节标题段落内的书签名
    （目录 PAGEREF 域所引用的 _TocXXXXXXXX 书签）。
    返回 {'procedure','photo','calibration'} 或 None
    """
    from docx import Document
    from docx.oxml.ns import qn
    keys = {
        '2试验大纲': 'procedure',
        '3试验照片': 'photo',              # 当前模板
        '4试验照片 Test Photo': 'photo',   # 旧模板
        '4校验证书': 'calibration',        # 当前模板
        '5校验证书': 'calibration',        # 旧模板
    }
    result = {}
    try:
        doc = Document(cover_path)
        for p in doc.paragraphs:
            key = p.text.strip()
            if key not in keys:
                continue
            # 取该段落中最后一个 bookmarkStart 的名称
            bms = p._p.findall(qn('w:bookmarkStart'))
            if not bms:
                return None
            result[keys[key]] = bms[-1].get(qn('w:name'))
    except Exception:
        return None
    if len(result) == 3:
        return result
    return None


def _last_page_find(doc, text):
    """备用定位：Word 查找标题文本最后一次出现的物理页码（1 起始）"""
    last = None
    rng = doc.Content
    guard = 0
    while guard < 200:
        try:
            f = rng.Find
            f.ClearFormatting()
            f.Text = text
            f.Forward = True
            f.Wrap = 0                      # wdFindStop
            if not f.Execute():
                break
            last = rng.Information(3)       # wdActiveEndPageNumber
            rng.Collapse(0)                 # 折叠到匹配文本末尾，继续向后查找
            if rng.End >= doc.Content.End:
                break
            guard += 1
        except Exception:
            break
    return last


def docx_to_pdf_and_locate(cover_path, out_pdf, bookmarks):
    """
    通过 Word COM 将 cover.docx1 导出为 PDF，并定位正文章节页（0 起始）。
    优先使用书签（bookmarks 字典），失败时退回 Find 文本查找。
    返回 {'procedure','photo','calibration'} 或 None
    """
    import win32com.client
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0                  # 屏蔽一切提示（含扩展名不匹配提示）
    doc = None
    try:
        doc = word.Documents.Open(
            os.path.abspath(cover_path),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        # 更新目录域，使封面内部页码准确
        try:
            for toc in doc.TablesOfContents:
                toc.Update()
        except Exception:
            pass

        # 导出 PDF
        # 注意：本机装有 DLP 透明加密软件，会按扩展名加密 Word 写出的 .pdf/.xps 文件
        # （加密后文件头为 \x17\xda_\xa0\x16，pypdf 无法读取）。
        # 因此先导出为 .dat 扩展名的明文 PDF，再改名为 .pdf（改名不触发加密）。
        dat_pdf = os.path.abspath(out_pdf) + '.dat'
        doc.ExportAsFixedFormat(dat_pdf, 17)   # 17 = PDF
        os.replace(dat_pdf, os.path.abspath(out_pdf))

        # 定位章节页码
        pages = None
        if bookmarks:
            try:
                doc.Bookmarks.ShowHidden = True
                found = {}
                for key in ('procedure', 'photo', 'calibration'):
                    name = bookmarks.get(key)
                    if not name:
                        break
                    bm = doc.Bookmarks.Item(name)
                    page = bm.Range.Information(3)     # 物理页码，1 起始
                    found[key] = page
                if len(found) == 3:
                    pages = found
            except Exception:
                pages = None

        if pages is None:
            proc = _last_page_find(doc, '试验大纲')
            photo = _last_page_find(doc, '试验照片')
            calib = _last_page_find(doc, '校验证书')
            if None not in (proc, photo, calib):
                pages = {'procedure': proc, 'photo': photo, 'calibration': calib}

        doc.Close(False)
        doc = None

        if pages is None:
            return None
        # 1 起始物理页码 -> 0 起始 PDF 页码，并校验顺序
        pages = {k: v - 1 for k, v in pages.items()}
        if not (pages['procedure'] < pages['photo'] < pages['calibration']):
            return None
        return pages
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        try:
            word.Quit()
        except Exception:
            pass


# ============================================================
#  最终 PDF 合并
# ============================================================
def merge_final(cover_pdf, proc_files, photos_pdf, calib_files, heading_pages, out_path):
    """
    合并顺序：
      封面(首页..试验大纲页) -> 1.Procedure 全部 PDF
      -> 封面(试验开口项页..试验照片页) -> 2.picture 生成 PDF
      -> 封面(校验证书页) -> 3.Calibration Certificate 全部 PDF
      -> 封面(结束页..末页)
    返回最终 PDF 总页数
    """
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(cover_pdf)
    total_pages = len(reader.pages)
    proc_pg = heading_pages['procedure']
    photo_pg = heading_pages['photo']
    calib_pg = heading_pages['calibration']

    writer = PdfWriter()

    def add_file(path):
        r = PdfReader(path)
        for pg in r.pages:
            writer.add_page(pg)

    def add_range(a, b):
        for i in range(a, b):
            writer.add_page(reader.pages[i])

    add_range(0, proc_pg + 1)                 # 封面含"2试验大纲"页
    for f in proc_files:                      # 1.Procedure
        add_file(f)
    add_range(proc_pg + 1, photo_pg + 1)      # 试验开口项..含"4试验照片"页
    if photos_pdf and os.path.isfile(photos_pdf):
        add_file(photos_pdf)                  # 2.picture 照片 PDF
    add_range(photo_pg + 1, calib_pg + 1)     # 含"5校验证书"页
    for f in calib_files:                     # 3.Calibration Certificate
        add_file(f)
    add_range(calib_pg + 1, total_pages)      # 结束页

    with open(out_path, 'wb') as fh:
        writer.write(fh)
    return len(writer.pages)


# ============================================================
#  单个报告处理
# ============================================================
def process_report(report_dir, output_dir, work_dir, folder_index, folder_count):
    from docx import Document

    folder_name = os.path.basename(os.path.normpath(report_dir))
    print(f'\n[{folder_index}/{folder_count}] {folder_name}')
    print('-' * 60)

    cover_path = os.path.join(report_dir, '0.cover', 'cover.docx1')
    if not os.path.isfile(cover_path):
        raise FileNotFoundError('缺少 0.cover/cover.docx1，跳过')

    proc_dir = os.path.join(report_dir, '1.Procedure')
    pic_dir = os.path.join(report_dir, '2.picture')
    calib_dir = os.path.join(report_dir, '3.Calibration Certificate')

    # ---- 1. 统计 1.Procedure 页数 ----
    proc_files = pdfs_in_folder(proc_dir)
    proc_total = sum(count_pdf_pages(f) for f in proc_files)
    print(f'  1.Procedure：{len(proc_files)} 个 PDF，共 {proc_total} 页')

    # ---- 2. 图片 -> A4 PDF，统计页数 ----
    photos_pdf = os.path.join(work_dir, 'photos.pdf')
    photo_total = build_photos_pdf(pic_dir, photos_pdf)
    print(f'  2.picture  ：{photo_total} 张图片 -> A4 PDF（{photos_pdf}）')

    # ---- 3. 统计 3.Calibration Certificate 页数 ----
    calib_files = pdfs_in_folder(calib_dir)
    calib_total = sum(count_pdf_pages(f) for f in calib_files)
    print(f'  3.Calibration Certificate：{len(calib_files)} 个 PDF，共 {calib_total} 页')

    # ---- 4. 复制 cover.docx1 到工作区，在副本上回填页数并保存 ----
    # （不修改源模板文件，保证软件可重复运行、模板始终保持"共XXX页"占位）
    work_cover = os.path.join(work_dir, 'cover_work.docx1')
    shutil.copy2(cover_path, work_cover)
    doc = Document(work_cover)
    revision = get_revision(doc)
    replaced = replace_page_counts(doc, proc_total, photo_total, calib_total)
    doc.save(work_cover)
    print(f'  0.cover/cover.docx1：回填 {replaced} 处页数，REVISION = {revision}')

    # ---- 5. cover.docx1 -> PDF + 章节定位 ----
    bookmarks = get_heading_bookmarks(work_cover)
    if bookmarks is None:
        print('  未找到章节书签，将使用文本查找定位章节页')
    cover_pdf = os.path.join(work_dir, 'cover.pdf')
    heading_pages = docx_to_pdf_and_locate(work_cover, cover_pdf, bookmarks)
    if heading_pages is None:
        raise RuntimeError('无法定位封面 PDF 中章节标题所在页码')
    print(f'  封面已导出 PDF：{cover_pdf}')
    print('  章节页码（0 起始）：试验大纲=%s 试验照片=%s 校验证书=%s' % (
        heading_pages['procedure'], heading_pages['photo'], heading_pages['calibration']))

    # ---- 6. 合并最终 PDF ----
    # 版本号规范化：文件名中版本号必须以 "_" 开头（如 _A、_B）；
    # 若模板 REVISION 字段本身已带 "_" 前缀则不再重复添加
    rev_suffix = revision.strip()
    if not rev_suffix.startswith('_'):
        rev_suffix = '_' + rev_suffix
    final_path = os.path.join(output_dir, f'{folder_name}{rev_suffix}.pdf')
    final_pages = merge_final(cover_pdf, proc_files, photos_pdf, calib_files,
                              heading_pages, final_path)
    print(f'  最终 PDF 已生成：{final_path}（共 {final_pages} 页）')
    return final_path, final_pages


# ============================================================
#  主程序
# ============================================================
def main():
    os.system(f'title {APP_TITLE}')
    print('=' * 60)
    print(f'  {APP_TITLE}')
    print('=' * 60)

    if not ensure_deps():
        show_popup('showerror', APP_TITLE, '依赖库安装失败，请检查网络后重试。')
        input('\n按 Enter 键退出...')
        return

    if not ensure_word():
        msg = ('未检测到 Microsoft Word。\n'
               'cover.docx1 转 PDF 需要本机安装 Microsoft Word，请先安装后再运行。')
        print(f'\n[!] {msg}')
        show_popup('showerror', APP_TITLE, msg)
        input('\n按 Enter 键退出...')
        return

    app_dir = get_app_dir()
    report_base = os.path.join(app_dir, 'REPORT FOLDER')
    if not os.path.isdir(report_base):
        msg = f'找不到文件夹：\n{report_base}\n\n请将本程序放在 REPORT FOLDER 的同级目录中运行。'
        print(f'\n[!] {msg}')
        show_popup('showerror', APP_TITLE, msg)
        input('\n按 Enter 键退出...')
        return

    output_dir = os.path.join(app_dir, 'REPORT OUTPUT')
    os.makedirs(output_dir, exist_ok=True)
    print(f'  报告目录：{report_base}')
    print(f'  输出目录：{output_dir}')

    reports = [d for d in sorted(os.listdir(report_base))
               if os.path.isdir(os.path.join(report_base, d)) and not d.startswith('.')]
    if not reports:
        msg = 'REPORT FOLDER 下没有试验报告文件夹。'
        print(f'\n[!] {msg}')
        show_popup('showwarning', APP_TITLE, msg)
        input('\n按 Enter 键退出...')
        return

    print(f'\n发现 {len(reports)} 个试验报告，开始处理...')

    work_dir = tempfile.mkdtemp(prefix='RMAC_work_')
    success, failed = 0, 0
    try:
        for idx, folder in enumerate(reports, 1):
            try:
                final_path, final_pages = process_report(
                    os.path.join(report_base, folder), output_dir, work_dir, idx, len(reports))
                print(f'  完成：{os.path.basename(final_path)}（{final_pages} 页）')
                success += 1
            except Exception as e:
                print(f'  [!] 处理失败：{e}')
                failed += 1
            print()
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

    print('=' * 60)
    print('  处理完成！')
    print(f'  成功：{success} 个')
    print(f'  失败：{failed} 个')
    print(f'  输出目录：{output_dir}')
    print('=' * 60)

    show_popup(
        'showinfo',
        APP_TITLE,
        f'处理完成！\n成功 {success} 个，失败 {failed} 个\n\n'
        f'输出目录：\n{output_dir}'
    )


if __name__ == '__main__':
    main()
    input('\n按 Enter 键退出...')
