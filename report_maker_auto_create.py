#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
试验报告生成器
=====================
功能：
  1. 选择试验大纲所在文件夹
  2. 解析文件名，提取试验编号和试验名称
  3. 将试验编号中的 FAT 改为 FAR、FTP 改为 FTR，形成试验报告编号
  4. 删除试验名称中的英文和特殊符号，只保留中文；
     若末尾为"试验大纲"或"试验"则循环去除，末尾加"试验报告"
     （避免"试验大纲试验报告"这类冗余名称）
  5. 根据报告编号和名称创建文件夹结构
     （含 0.cover / 1.Procedure / 2.picture / 3.Calibration Certificate）
  6. 使用 baseline_cover.docx1 模板生成封面文档：
     - 先命名为 cover，生成结束后改后缀为 .docx1
     - 填入 项目名称(=试验大纲所在文件夹名)、文件名称(=报告名称)、
       文件编号(=报告编号)、文件版本(A)
     - 同步更新页脚尾标中的 文件名称 / 文件编号 / 版本
  7. 所有生成物放置在 REPORT FOLDER 文件夹中
"""

import os
import re
import sys
import shutil

# tkinter 延迟到实际使用时再导入（保证模块在无 tkinter 环境下也可被导入/测试）
# Word 文档 XML 命名空间
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# 窗口标题
APP_TITLE = 'REPORT MAKER AUTO CREATE v1.0 by bohangyang 2026.08'

# python-docx 在模块级导入，便于函数直接使用；缺失时置 None，由 __main__ 负责安装
try:
    from docx import Document
except ImportError:
    Document = None


# ============================================================
#  文件夹选择
# ============================================================
def select_folder(title="请选择文件夹"):
    """弹出系统文件夹选择对话框，返回选中路径或空字符串"""
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.title(APP_TITLE)
    root.withdraw()
    root.attributes('-topmost', True)
    folder = filedialog.askdirectory(title=title)
    root.destroy()
    return folder


# ============================================================
#  文件名解析
# ============================================================
def parse_filename(filename):
    """
    解析文件名，提取试验编号和试验名称。
    文件名前半段为试验编号（如 PB1Y-VEH-FAT-16-0001），后半段为试验名称。
    """
    name, ext = os.path.splitext(filename)
    if not ext and '.' in filename:
        name = filename

    # 正则匹配试验编号：前缀-前缀-三位字母-两位数字-四位数字
    pattern = r'^([A-Za-z0-9]+(?:-[A-Za-z0-9]+)*-[A-Z]{3}-\d{2}-\d{4})'
    match = re.match(pattern, name)

    if match:
        test_number = match.group(1)
        rest = name[match.end():]
    else:
        # 回退方案：以第一个中文字符为分界
        cjk_match = re.search(r'[\u4e00-\u9fff]', name)
        if cjk_match:
            test_number = name[:cjk_match.start()].rstrip('-_ ')
            rest = name[cjk_match.start():]
        else:
            test_number = name
            rest = ''

    return test_number, rest


# ============================================================
#  编号 / 名称处理
# ============================================================
def process_test_number(test_number):
    """
    将试验编号中的 FAT 改为 FAR、FTP 改为 FTR，形成试验报告编号。
    """
    return test_number.replace('FAT', 'FAR').replace('FTP', 'FTR')


def process_test_name(test_name):
    """
    删除试验名称中的英文和特殊符号，只保留中文。
    若末尾为"试验大纲"或"试验"则循环去除，末尾加"试验报告"。
    避免"试验大纲试验报告"这类冗余名称。
    """
    name = ''.join(re.findall(r'[\u4e00-\u9fff]', test_name))
    # 循环去除末尾的"试验大纲"或"试验"（先匹配较长的"试验大纲"）
    while True:
        if name.endswith('试验大纲'):
            name = name[:-4]
        elif name.endswith('试验'):
            name = name[:-2]
        else:
            break
    return name + '试验报告'


# ============================================================
#  Word 文档编辑（SDT 内容控件）
# ============================================================
def set_sdt_text(sdt_element, new_text):
    """
    设置 SDT（结构化文档标签）内容控件的文本。
    移除占位符标记，替换文本内容，并将 SDT 的 sdtPr/rPr 格式属性
    复制到 run 的 rPr，确保字体大小等格式正确继承（而非回退到段落样式）。
    """
    from lxml import etree
    from copy import deepcopy

    # 移除 showingPlcHdr（占位符标记）
    sdtPr = sdt_element.find(f'{{{W_NS}}}sdtPr')
    if sdtPr is not None:
        plcHdr = sdtPr.find(f'{{{W_NS}}}showingPlcHdr')
        if plcHdr is not None:
            sdtPr.remove(plcHdr)

    # 获取 SDT 自身的 rPr（定义了正确的字号、字体等格式）
    sdt_rpr = sdtPr.find(f'{{{W_NS}}}rPr') if sdtPr is not None else None

    # 查找 sdtContent
    sdtContent = sdt_element.find(f'{{{W_NS}}}sdtContent')
    if sdtContent is None:
        return

    # 查找所有段落，只保留第一个
    paragraphs = sdtContent.findall(f'.//{{{W_NS}}}p')
    if not paragraphs:
        return
    for p in paragraphs[1:]:
        p.getparent().remove(p)

    first_p = paragraphs[0]

    # 在段落中只保留第一个 run
    runs = first_p.findall(f'{{{W_NS}}}r')
    if runs:
        for r in runs[1:]:
            first_p.remove(r)

        # 用 SDT 的 rPr 替换 run 的 rPr，确保格式（字号等）正确继承
        # 原因：run 的 rPr 可能只有占位符样式引用（rStyle=a8），
        #        没有显式字号；移除 showingPlcHdr 后会回退到段落样式
        #        的字号（如 FieldText 样式的 sz=28/14pt），导致字体变大。
        run_rpr = runs[0].find(f'{{{W_NS}}}rPr')
        if run_rpr is not None:
            runs[0].remove(run_rpr)
        if sdt_rpr is not None:
            runs[0].insert(0, deepcopy(sdt_rpr))

        t_elements = runs[0].findall(f'{{{W_NS}}}t')
        if t_elements:
            t_elements[0].text = new_text
            for t in t_elements[1:]:
                runs[0].remove(t)
        else:
            t = etree.SubElement(runs[0], f'{{{W_NS}}}t')
            t.text = new_text
    else:
        r = etree.SubElement(first_p, f'{{{W_NS}}}r')
        if sdt_rpr is not None:
            r.insert(0, deepcopy(sdt_rpr))
        t = etree.SubElement(r, f'{{{W_NS}}}t')
        t.text = new_text


def find_table_sdts(table):
    """
    查找表格中每行的 SDT 元素。
    返回列表，每个元素是该行的 SDT 列表。
    """
    sdts_by_row = []
    for tr in table._tbl.findall(f'{{{W_NS}}}tr'):
        row_sdts = tr.findall(f'.//{{{W_NS}}}sdt')
        sdts_by_row.append(row_sdts)
    return sdts_by_row


def _set_cell_text(cell, new_text):
    """
    设置单元格文本，保留首个 run 的格式，清除其余 run 与多余段落。
    用于页脚中"文件名称"这类普通文本单元格。
    """
    if not cell.paragraphs:
        return
    # 只保留第一个段落
    first_p = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # 在首段中保留首个 run 并替换文本
    runs = first_p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        first_p.add_run(new_text)


def edit_footer(doc, report_name, report_number, file_version):
    """
    更新页脚（尾标）表格：
      - 文件名称（普通文本单元格）-> 报告名称
      - 文件编号（SDT，绑定 keywords）-> 报告编号
      - 版本（SDT，绑定 subject）-> 文件版本
    文档可能有多个节，后续节页脚通常链接到首节；此处遍历所有节以防万一。
    """
    for section in doc.sections:
        footer = section.footer
        if not footer.tables:
            continue
        table = footer.tables[0]
        if len(table.rows) < 2 or len(table.columns) < 3:
            continue
        # 第2行中间列为"文件名称"值（普通文本，模板中为"标准化APM例行试验报告模板"）
        _set_cell_text(table.rows[1].cells[1], report_name)
        # 页脚中的 SDT：第1个=文件编号(Code/keywords)，第2个=版本(Version/subject)
        sdts = footer._element.findall(f'.//{{{W_NS}}}sdt')
        if len(sdts) >= 1:
            set_sdt_text(sdts[0], report_number)
        if len(sdts) >= 2:
            set_sdt_text(sdts[1], file_version)


def edit_cover_document(template_path, output_path, project_name,
                        report_name, report_number, file_version='A'):
    """
    复制模板并编辑封面文档：
      封面表格: 项目名称 / 文件名称(报告名称) / 文件编号(报告编号) / 文件版本(A)
      页脚尾标: 文件名称 / 文件编号 / 版本 同步更新
      核心属性: 同步写入（确保数据绑定一致）
    """
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # 封面信息表 = tables[1]
    table = doc.tables[1]
    sdts_by_row = find_table_sdts(table)
    values = [
        project_name,    # 项目名称
        report_name,     # 文件名称 = 报告名称
        report_number,   # 文件编号 = 报告编号
        file_version,    # 文件版本
    ]
    for i, value in enumerate(values):
        if i < len(sdts_by_row) and sdts_by_row[i]:
            set_sdt_text(sdts_by_row[i][0], value)

    # 页脚尾标
    edit_footer(doc, report_name, report_number, file_version)

    # 同步核心属性（确保 SDT 数据绑定一致）
    cp = doc.core_properties
    cp.comments = project_name    # dc:description -> 项目名称
    cp.title = report_name        # dc:title -> 文件名称
    cp.keywords = report_number   # cp:keywords -> 文件编号
    cp.subject = file_version     # dc:subject -> 文件版本

    doc.save(output_path)


# ============================================================
#  文件夹结构创建
# ============================================================
def create_report_structure(base_path, report_number, report_name):
    """
    在 base_path 下创建：
      报告编号_报告名称/
        0.cover/
        1.Procedure/
        2.picture/
        3.Calibration Certificate/
    """
    folder_name = f"{report_number}_{report_name}"
    main_folder = os.path.join(base_path, folder_name)

    subfolders = [
        '0.cover',
        '1.Procedure',
        '2.picture',
        '3.Calibration Certificate',
    ]
    for sub in subfolders:
        os.makedirs(os.path.join(main_folder, sub), exist_ok=True)

    return main_folder


# ============================================================
#  应用程序目录定位
# ============================================================
def get_app_dir():
    """
    获取应用程序所在目录。
    - 打包为 exe 时（PyInstaller），sys.executable 指向 exe 文件路径，
      返回 exe 所在目录（而非 _MEIPASS 临时解压目录）。
    - 以 .py 脚本运行时，返回脚本所在目录。
    """
    if getattr(sys, 'frozen', False):
        # PyInstaller 打包后，sys.executable = exe 文件完整路径
        return os.path.dirname(os.path.abspath(sys.executable))
    else:
        return os.path.dirname(os.path.abspath(__file__))


# ============================================================
#  模板文件查找
# ============================================================
def find_template(app_dir, desktop_path):
    """查找 baseline_cover.docx1 模板文件：优先程序目录(exe同目录)，其次桌面"""
    for path in [
        os.path.join(app_dir, 'baseline_cover.docx1'),
        os.path.join(desktop_path, 'baseline_cover.docx1'),
    ]:
        if os.path.exists(path):
            return path
    return None


# ============================================================
#  主程序
# ============================================================
def main():
    from tkinter import messagebox
    # 设置控制台窗口标题
    os.system(f'title {APP_TITLE}')
    print("=" * 60)
    print(f"  {APP_TITLE}")
    print("=" * 60)

    script_dir = get_app_dir()
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')

    # ---- 查找模板 ----
    template_path = find_template(script_dir, desktop_path)
    if template_path is None:
        msg = "找不到模板文件 baseline_cover.docx1\n请将其放在 exe 同目录或桌面中。"
        print(f"\n[!] {msg}")
        messagebox.showerror(APP_TITLE, msg)
        input("\n按 Enter 键退出...")
        return

    print(f"  模板文件：{template_path}")

    # ---- 选择试验大纲文件夹 ----
    print("\n>>> 请选择试验大纲所在的文件夹...")
    source_folder = select_folder("请选择试验大纲所在的文件夹")
    if not source_folder:
        print("未选择文件夹，程序退出。")
        return

    print(f"  选中文件夹：{source_folder}")

    # 项目名称 = 试验大纲所在目录的文件夹名称
    project_name = os.path.basename(os.path.normpath(source_folder))
    print(f"  项目名称：{project_name}")

    # ---- 创建 REPORT FOLDER ----
    report_base = os.path.join(script_dir, 'REPORT FOLDER')
    os.makedirs(report_base, exist_ok=True)
    print(f"  报告输出目录：{report_base}")

    # ---- 列出文件 ----
    skip_names = {'desktop.ini', 'Thumbs.db', '.DS_Store'}
    files = [f for f in os.listdir(source_folder)
             if os.path.isfile(os.path.join(source_folder, f))
             and f not in skip_names]

    if not files:
        print("\n[!] 选中的文件夹中没有文件。")
        input("按 Enter 键退出...")
        return

    print(f"\n找到 {len(files)} 个文件，开始处理...\n")
    print("-" * 60)

    success_count = 0
    error_count = 0

    for idx, filename in enumerate(files, 1):
        print(f"[{idx}/{len(files)}] {filename}")

        # 解析文件名
        test_number, test_name = parse_filename(filename)

        if not test_number:
            print(f"  [!] 无法解析试验编号，跳过")
            error_count += 1
            continue

        # 处理编号和名称
        report_number = process_test_number(test_number)
        report_name = process_test_name(test_name)

        print(f"  试验编号：{test_number}")
        print(f"  试验名称：{test_name}")
        print(f"  报告编号：{report_number}")
        print(f"  报告名称：{report_name}")

        # 创建文件夹结构
        main_folder = create_report_structure(report_base, report_number, report_name)
        print(f"  文件夹已创建：{os.path.basename(main_folder)}")

        # 生成封面文档：先命名为 cover，生成结束后改后缀为 .docx1
        cover_dir = os.path.join(main_folder, '0.cover')
        temp_cover = os.path.join(cover_dir, 'cover')
        final_cover = os.path.join(cover_dir, 'cover.docx1')
        try:
            edit_cover_document(
                template_path, temp_cover,
                project_name, report_name, report_number, 'A'
            )
            # 生成结束后修改后缀名为 .docx1
            if os.path.exists(final_cover):
                os.remove(final_cover)
            os.rename(temp_cover, final_cover)
            print(f"  封面文档已生成：cover.docx1")
            success_count += 1
        except Exception as e:
            print(f"  [!] 封面文档生成失败：{e}")
            error_count += 1

        print()

    # ---- 汇总 ----
    print("=" * 60)
    print(f"  处理完成！")
    print(f"  成功：{success_count} 个")
    print(f"  失败：{error_count} 个")
    print(f"  报告目录：{report_base}")
    print("=" * 60)

    messagebox.showinfo(
        APP_TITLE,
        f"成功创建 {success_count} 个试验报告\n"
        f"失败：{error_count} 个\n\n"
        f"报告目录：\n{report_base}"
    )


if __name__ == '__main__':
    if Document is None:
        print("[!] 缺少依赖库 python-docx，正在尝试安装...")
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
        try:
            from docx import Document
        except ImportError:
            print("安装失败，请手动执行: pip install python-docx")
            input("按 Enter 键退出...")
            sys.exit(0)

    main()
    input("\n按 Enter 键退出...")
